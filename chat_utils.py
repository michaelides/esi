import os
import json
import uuid
import re
from typing import List, Dict, Any, Tuple, Optional # Added Optional
import streamlit as st
from llama_index.core.llms import ChatMessage, MessageRole

# Configuration imports
from config import MEMORY_DIR, MAX_CHAT_HISTORY_MESSAGES, DEFAULT_PROMPTS

# Agent-related imports
from agent import generate_llm_greeting, generate_suggested_prompts

# Session management imports
from session_manager import save_chat_history, save_chat_metadata

# Document generation imports
from docx import Document
from io import BytesIO

# --- Functions moved from app.py ---

@st.cache_data(show_spinner=False)
def get_initial_greeting_text() -> str:
    """Generates and caches the initial LLM greeting text for startup."""
    print("LOG: chat_utils.get_initial_greeting_text() CALLED")
    try:
        return generate_llm_greeting()
    except Exception as e:
        print(f"ERROR: chat_utils.get_initial_greeting_text: Failed to generate LLM greeting: {e}")
        return "Hello! I'm ESI, your AI assistant. How can I help you today? (Default Greeting)"

@st.cache_data(show_spinner=False)
def cached_generate_suggested_prompts(chat_history: List[Dict[str, Any]]) -> List[str]:
    """
    Generates suggested prompts based on chat history, cached to avoid redundant LLM calls.
    """
    print("LOG: chat_utils.cached_generate_suggested_prompts() CALLED")
    try:
        return generate_suggested_prompts(chat_history)
    except Exception as e:
        print(f"ERROR: chat_utils.cached_generate_suggested_prompts: Failed to generate suggested prompts: {e}")
        return DEFAULT_PROMPTS[:] # Return a copy of default prompts

def format_chat_history(streamlit_messages: List[Dict[str, Any]]) -> List[ChatMessage]:
    """
    Converts Streamlit message history to LlamaIndex ChatMessage list.
    """
    if not isinstance(streamlit_messages, list):
        print(f"ERROR: chat_utils.format_chat_history: Expected a list for streamlit_messages, got {type(streamlit_messages)}. Returning empty list.")
        return []

    truncated_messages = streamlit_messages[-MAX_CHAT_HISTORY_MESSAGES:]
    history = []
    for msg in truncated_messages:
        if not isinstance(msg, dict) or "role" not in msg or "content" not in msg:
            print(f"WARNING: chat_utils.format_chat_history: Skipping invalid message format: {msg}")
            continue
        role = MessageRole.USER if msg["role"] == "user" else MessageRole.ASSISTANT
        history.append(ChatMessage(role=role, content=str(msg["content"]))) # Ensure content is string
    return history

def create_new_chat_session_in_memory():
    """
    Creates a new chat session in memory and sets it as the current chat.
    """
    try:
        new_chat_id = str(uuid.uuid4())
        new_chat_name = "Current Session"
        if st.session_state.get("long_term_memory_enabled", False): # Safely access LTM flag
            existing_idea_nums = []
            for name in st.session_state.get("chat_metadata", {}).values(): # Safely access chat_metadata
                if isinstance(name, str): # Ensure name is a string before regex
                    match = re.match(r"Idea (\d+)", name)
                    if match:
                        existing_idea_nums.append(int(match.group(1)))
            next_idea_num = max(existing_idea_nums) + 1 if existing_idea_nums else 1
            new_chat_name = f"Idea {next_idea_num}"

        st.session_state.chat_metadata[new_chat_id] = new_chat_name
        st.session_state.all_chat_messages[new_chat_id] = [{"role": "assistant", "content": get_initial_greeting_text()}]
        st.session_state.current_chat_id = new_chat_id
        st.session_state.messages = st.session_state.all_chat_messages[new_chat_id]
        st.session_state.chat_modified = False
        st.session_state.suggested_prompts = cached_generate_suggested_prompts(st.session_state.messages)
        print(f"LOG: chat_utils.Created new chat in memory: ID={new_chat_id}, Name='{new_chat_name}'")
        return new_chat_id
    except Exception as e:
        print(f"ERROR: chat_utils.create_new_chat_session_in_memory: Failed to create new chat. Error: {e}")
        # Attempt to set a very basic fallback state if possible
        st.session_state.messages = [{"role": "assistant", "content": "Error: Could not start a new chat session."}]
        st.session_state.suggested_prompts = DEFAULT_PROMPTS[:]
        return None # Indicate failure

def switch_chat(chat_id: str):
    """Switches to an existing chat, lazy-loading messages if necessary."""
    if not st.session_state.get("long_term_memory_enabled", False):
        print("LOG: chat_utils.LTM disabled. Cannot switch. Creating new temporary session.")
        create_new_chat_session_in_memory() # This might rerun
        st.rerun() # Ensure rerun if create_new_chat_session_in_memory doesn't always
        return

    if chat_id not in st.session_state.get("chat_metadata", {}):
        print(f"ERROR: chat_utils.switch_chat: Attempted to switch to chat ID '{chat_id}' not found in metadata.")
        st.error(f"Could not switch to chat '{chat_id}'. Chat not found.")
        return

    if st.session_state.all_chat_messages.get(chat_id) is None: # Lazy load
        print(f"LOG: chat_utils.Messages for chat ID '{chat_id}' not loaded. Loading from disk...")
        user_id = st.session_state.get("user_id")
        if not user_id:
            print(f"ERROR: chat_utils.switch_chat: No user_id in session_state. Cannot load chat '{chat_id}'.")
            st.error("User session error. Cannot load chat.")
            return

        user_dir = os.path.join(MEMORY_DIR, user_id)
        chat_file = os.path.join(user_dir, f"{chat_id}.json")

        if os.path.exists(chat_file):
            try:
                with open(chat_file, "r", encoding="utf-8") as f:
                    st.session_state.all_chat_messages[chat_id] = json.load(f)
                print(f"LOG: chat_utils.Successfully loaded messages for chat ID '{chat_id}'.")
            except (json.JSONDecodeError, IOError, OSError) as e:
                print(f"ERROR: chat_utils.switch_chat: Error loading/decoding chat file {chat_file}: {e}.")
                st.error(f"Failed to load chat '{st.session_state.chat_metadata.get(chat_id, chat_id)}'. Content may be corrupted.")
                st.session_state.all_chat_messages[chat_id] = [{"role": "assistant", "content": f"Error: Could not load chat content for '{st.session_state.chat_metadata.get(chat_id, chat_id)}'."}]
        else:
            print(f"LOG: chat_utils.Chat file {chat_file} not found for chat ID '{chat_id}'. Setting to empty messages (error state).")
            st.error(f"Chat file for '{st.session_state.chat_metadata.get(chat_id, chat_id)}' not found.")
            st.session_state.all_chat_messages[chat_id] = [{"role": "assistant", "content": f"Error: Chat file not found for '{st.session_state.chat_metadata.get(chat_id, chat_id)}'."}]

    st.session_state.current_chat_id = chat_id
    st.session_state.messages = st.session_state.all_chat_messages.get(chat_id, []) # Default to empty list if still None
    if st.session_state.messages is None: # Should be handled by above, but safeguard
        st.session_state.messages = []
        st.session_state.all_chat_messages[chat_id] = []

    st.session_state.suggested_prompts = cached_generate_suggested_prompts(st.session_state.messages)
    st.session_state.chat_modified = True
    print(f"LOG: chat_utils.Switched to chat: ID={chat_id}, Name='{st.session_state.chat_metadata.get(chat_id, 'Unknown')}'")
    st.rerun()

def delete_chat_session(chat_id: str):
    """Deletes a chat history and its metadata."""
    if not st.session_state.get("long_term_memory_enabled", False):
        print("LOG: chat_utils.LTM disabled. Cannot delete. Resetting current session.")
        if chat_id == st.session_state.get("current_chat_id"):
            create_new_chat_session_in_memory()
        st.rerun()
        return

    user_id = st.session_state.get("user_id")
    if not user_id:
        print("ERROR: chat_utils.delete_chat_session: No user_id. Cannot delete chat.")
        st.error("User session error. Cannot delete chat.")
        return

    is_current_chat = (chat_id == st.session_state.get("current_chat_id"))

    if chat_id in st.session_state.get("all_chat_messages", {}):
        del st.session_state.all_chat_messages[chat_id]
    if chat_id in st.session_state.get("chat_metadata", {}):
        del st.session_state.chat_metadata[chat_id]

    user_dir = os.path.join(MEMORY_DIR, user_id)
    chat_file = os.path.join(user_dir, f"{chat_id}.json")
    try:
        if os.path.exists(chat_file):
            os.remove(chat_file)
            print(f"LOG: chat_utils.Deleted chat file: {chat_file}")
    except OSError as e:
        print(f"ERROR: chat_utils.delete_chat_session: Error deleting chat file {chat_file}: {e}")
        st.error(f"Could not delete chat file for '{chat_id}'.")
        # Proceed to save metadata anyway, as the chat is removed from memory.

    save_chat_metadata(user_id, st.session_state.chat_metadata) # save_chat_metadata has its own error handling
    print(f"LOG: chat_utils.Deleted chat from memory: ID={chat_id}. Updated metadata saved.")

    if is_current_chat:
        if st.session_state.chat_metadata:
            first_available_chat_id = next(iter(st.session_state.chat_metadata))
            print(f"LOG: chat_utils.Deleted current chat. Switching to: {first_available_chat_id}")
            switch_chat(first_available_chat_id)
        else:
            print("LOG: chat_utils.Deleted last chat. Creating new session.")
            create_new_chat_session_in_memory() # This will call st.rerun()
    else: # If a non-current chat was deleted, just rerun to update UI
        st.rerun()


def rename_chat(chat_id: str, new_name: str):
    """Renames the specified chat."""
    if not st.session_state.get("long_term_memory_enabled", False):
        print("LOG: chat_utils.LTM disabled. Cannot rename chats.")
        return

    user_id = st.session_state.get("user_id")
    if not user_id:
        print(f"ERROR: chat_utils.rename_chat: No user_id. Cannot rename chat '{chat_id}'.")
        return

    if chat_id and new_name and st.session_state.get("chat_metadata", {}).get(chat_id) != new_name:
        st.session_state.chat_metadata[chat_id] = new_name
        save_chat_metadata(user_id, st.session_state.chat_metadata) # save_chat_metadata handles its errors
        print(f"LOG: chat_utils.Renamed chat {chat_id} to '{new_name}'")
    # No st.rerun() here, handled by Streamlit's on_change for text_input typically

def get_discussion_markdown(chat_id: str) -> str:
    """Retrieves messages for a given chat_id and formats them into a Markdown string."""
    messages = st.session_state.get("all_chat_messages", {}).get(chat_id, [])
    if not messages:
        return "No messages found for this chat."
    markdown_content = []
    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        markdown_content.append(f"**{role}:**\n{content}\n\n---")
    return "\n".join(markdown_content)

def get_discussion_docx(chat_id: str) -> bytes:
    """Retrieves messages for a given chat_id and formats them into a DOCX file."""
    messages = st.session_state.get("all_chat_messages", {}).get(chat_id, [])
    document = Document()

    chat_name = st.session_state.get("chat_metadata", {}).get(chat_id, "Untitled Chat")
    document.add_heading(f"Chat Discussion: {chat_name}", level=1)
    # document.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}") # Consider adding timestamp

    if not messages:
        document.add_paragraph("No messages found for this chat.")

    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        try:
            document.add_heading(f"{role}:", level=3)
            document.add_paragraph(content)
            document.add_paragraph("---")
        except Exception as e:
            print(f"ERROR: chat_utils.get_discussion_docx: Error adding content to DOCX for chat {chat_id}. Role: {role}. Error: {e}")
            document.add_paragraph(f"Error processing message for role {role}. Content may be incomplete.")

    byte_stream = BytesIO()
    try:
        document.save(byte_stream)
    except Exception as e:
        print(f"ERROR: chat_utils.get_discussion_docx: Failed to save DOCX to BytesIO for chat {chat_id}. Error: {e}")
        # Return a simple DOCX with error message
        error_doc = Document()
        error_doc.add_paragraph(f"Failed to generate DOCX export for chat {chat_name} due to an error: {e}")
        error_byte_stream = BytesIO()
        error_doc.save(error_byte_stream)
        error_byte_stream.seek(0)
        return error_byte_stream.getvalue()

    byte_stream.seek(0)
    return byte_stream.getvalue()

print("DEBUG: chat_utils.py processed with error handling improvements.")
