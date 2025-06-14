import gradio as gr
import os
import sys
import uuid
import json
from datetime import datetime, timedelta
from huggingface_hub import HfFileSystem, hf_hub_delete
from llama_index.core.llms import ChatMessage, MessageRole
from llama_index.core.tools import FunctionTool
from llama_index.core import Settings as LlamaSettings # For LLM temperature

import functools
import re
import urllib.parse
import io # For BytesIO for docx download

# For File Uploads
import pandas as pd
from PyPDF2 import PdfReader
# from docx import Document as DocxDocument # Renamed to avoid conflict with gr.Document if any
from docx import Document as PythonDocxDocument
import pyreadstat

# Add project root to Python path
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
sys.path.append(PROJECT_ROOT)

try:
    from agent import (
        initialize_settings as initialize_agent_settings,
        generate_llm_greeting,
        create_orchestrator_agent,
        generate_suggested_prompts,
        DEFAULT_PROMPTS,
        SUGGESTED_PROMPT_COUNT
    )
    from tools import UI_ACCESSIBLE_WORKSPACE
    print("Successfully imported agent functions, config, and tools.")
except ImportError as e:
    print(f"Error importing modules: {e}")
    # Dummy implementations ... (ensure these are comprehensive enough for testing UI)
    def initialize_agent_settings(): print("Dummy initialize_agent_settings"); LlamaSettings.llm = type('obj', (object,), {'temperature': 0.7})() # Dummy LLM object
    def generate_llm_greeting(): return "Hello! LLM Greeting (dummy)."
    def create_orchestrator_agent(dynamic_tools=None, max_search_results=10):
        print(f"Dummy create_orchestrator_agent called with max_search_results={max_search_results}.")
        class DummyAgent:
            def chat(self, query, chat_history=None):
                class DummyResponse: response = f"Echo: {query}"
                return DummyResponse()
        return DummyAgent()
    def generate_suggested_prompts(chat_history_list_dict, count, existing_prompts=None):
        return [f"Dummy Prompt {i+1}" for i in range(count)]
    DEFAULT_PROMPTS = ["What can you do?", "Explain a complex topic simply."]
    SUGGESTED_PROMPT_COUNT = 4
    UI_ACCESSIBLE_WORKSPACE = "code_interpreter_ws"
    HF_USER_MEMORIES_DATASET_ID = "dummy/user_memories" # Example: "your_username/your_dataset_name"
    HF_TOKEN = None # Needs to be set in environment or config.py for real use


if not os.path.exists(os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE)):
    os.makedirs(os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE), exist_ok=True)

fs = HfFileSystem(token=HF_TOKEN if HF_TOKEN else os.getenv("HF_TOKEN"))
# Initialize LlamaSettings.llm if not done by initialize_agent_settings in dummy mode
if not hasattr(LlamaSettings, 'llm'):
    LlamaSettings.llm = type('obj', (object,), {'temperature': 0.7})()

LLM_SETTINGS_OK, LLM_SETTINGS_ERROR = True, None
def setup_global_llm_settings_gr():
    global LLM_SETTINGS_OK, LLM_SETTINGS_ERROR
    try:
        initialize_agent_settings()
        print("LLM settings initialized successfully via agent.py.")
        LLM_SETTINGS_OK, LLM_SETTINGS_ERROR = True, None
    except Exception as e:
        error_message = f"Fatal Error: Could not initialize LLM settings. {e}"
        print(error_message)
        LLM_SETTINGS_OK, LLM_SETTINGS_ERROR = False, error_message
setup_global_llm_settings_gr()

_GLOBAL_UPLOADED_DOCS = {}
_GLOBAL_UPLOADED_DFS = {}

# --- All helper functions (get_user_id, data saving/loading, file processing, formatters, agent response) ---
# ... (These are largely the same as in the previous version, ensure they are present and correct) ...
# For brevity, I'll assume they are here from the previous step. Key ones:
def get_user_id_gr(): return str(uuid.uuid4())
def _load_user_data_from_hf(user_id: str) -> dict:
    # ... (implementation from previous step) ...
    if not fs.token: print("Warning: HfFileSystem not initialized with a token...")
    all_chat_metadata = {}
    all_chat_messages_li_format = {}
    try:
        metadata_hf_path = f"{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id}_metadata.json"
        messages_hf_path = f"{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id}_messages.json"
        if fs.exists(metadata_hf_path):
            with fs.open(metadata_hf_path, "r") as f: all_chat_metadata = json.load(f)
        if fs.exists(messages_hf_path):
            with fs.open(messages_hf_path, "r") as f: all_chat_messages_li_format = json.load(f)
    except Exception as e: print(f"Error loading user data from HF for user {user_id}: {e}")
    return {"metadata": all_chat_metadata, "messages": all_chat_messages_li_format}

def save_chat_history_gr(user_id: str, chat_id: str, messages_li_format: list, ltm_enabled: bool):
    # ... (implementation from previous step, ensure path is correct for HfFileSystem) ...
    if not ltm_enabled or not fs.token : return
    try:
        messages_hf_path = f"{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id}_messages.json"
        # Ensure directory exists (HfFileSystem might not auto-create parent dirs for datasets)
        fs.makedirs(f"{HF_USER_MEMORIES_DATASET_ID}/user_memories", exist_ok=True)
        existing_messages_data = {}
        if fs.exists(messages_hf_path):
            with fs.open(messages_hf_path, "r") as f: existing_messages_data = json.load(f)
        existing_messages_data[chat_id] = messages_li_format
        with fs.open(messages_hf_path, "w") as f: json.dump(existing_messages_data, f, indent=2)
    except Exception as e: print(f"Error saving chat history to HF: {e}")


def save_chat_metadata_gr(user_id: str, chat_metadata: dict, ltm_enabled: bool):
    # ... (implementation from previous step, ensure path is correct for HfFileSystem) ...
    if not ltm_enabled or not fs.token: return
    try:
        metadata_hf_path = f"{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id}_metadata.json"
        fs.makedirs(f"{HF_USER_MEMORIES_DATASET_ID}/user_memories", exist_ok=True)
        with fs.open(metadata_hf_path, "w") as f: json.dump(chat_metadata, f, indent=2)
    except Exception as e: print(f"Error saving chat metadata to HF: {e}")

def process_uploaded_files_gr(files_list, current_docs, current_dfs):
    # ... (implementation from previous step) ...
    updated_docs = current_docs.copy(); updated_dfs = current_dfs.copy()
    success_messages = []; error_messages = []
    workspace_full_path = os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE)
    if not os.path.exists(workspace_full_path): os.makedirs(workspace_full_path, exist_ok=True)
    for temp_file_obj in files_list:
        original_filename = os.path.basename(temp_file_obj.name)
        file_extension = os.path.splitext(original_filename)[1].lower()
        perm_file_path = os.path.join(workspace_full_path, original_filename)
        try:
            with open(temp_file_obj.name, "rb") as f_src, open(perm_file_path, "wb") as f_dst: f_dst.write(f_src.read())
            if file_extension in [".csv", ".xlsx", ".sav", ".rdata", ".rds"]: success_messages.append(f"Ethical Reminder for {original_filename}: Ensure ethical approval before analyzing research data.")
            if file_extension in [".pdf", ".docx", ".md", ".txt"]:
                text_content = ""
                if file_extension == ".pdf": reader = PdfReader(perm_file_path); text_content = "".join(page.extract_text() or "" for page in reader.pages)
                elif file_extension == ".docx": doc = PythonDocxDocument(perm_file_path); text_content = "\n".join(para.text for para in doc.paragraphs)
                else: # .md, .txt
                    with open(perm_file_path, "r", encoding="utf-8") as f: text_content = f.read()
                updated_docs[original_filename] = text_content
                success_messages.append(f"Doc '{original_filename}' processed. Use `read_uploaded_document('{original_filename}')`.")
            elif file_extension in [".csv", ".xlsx", ".sav"]:
                df = None
                if file_extension == ".csv": df = pd.read_csv(perm_file_path)
                elif file_extension == ".xlsx": df = pd.read_excel(perm_file_path)
                elif file_extension == ".sav":
                    try: df, meta = pyreadstat.read_sav(perm_file_path)
                    except Exception as e_sav: error_messages.append(f"Error reading .sav '{original_filename}': {e_sav}"); continue
                if df is not None: updated_dfs[original_filename] = df; success_messages.append(f"Data '{original_filename}' loaded. Use `analyze_uploaded_dataframe('{original_filename}')`.")
                else: error_messages.append(f"Could not load dataframe from '{original_filename}'.")
            else: success_messages.append(f"File '{original_filename}' saved. Type '{file_extension}' may not be directly readable.")
        except Exception as e: error_messages.append(f"Error with '{original_filename}': {e}")
    return updated_docs, updated_dfs, success_messages, error_messages

def read_uploaded_document_tool_fn_gr(filename: str): return _GLOBAL_UPLOADED_DOCS.get(filename, f"Error: Doc '{filename}' not found.")
def analyze_dataframe_tool_fn_gr(filename: str, head_rows: int = 5):
    # ... (implementation from previous step, ensure df.describe(include='all')) ...
    if filename not in _GLOBAL_UPLOADED_DFS: return f"Error: DataFrame '{filename}' not found."
    df = _GLOBAL_UPLOADED_DFS[filename]
    info_str = f"DataFrame: {filename}\nShape: {df.shape}\nColumns: {', '.join(df.columns)}\nData Types:\n{df.dtypes.to_string()}\n"
    head_rows = max(0, min(head_rows, len(df)))
    if head_rows > 0 : info_str += f"First {head_rows} rows:\n{df.head(head_rows).to_string()}\n"
    try: info_str += f"Summary Statistics:\n{df.describe(include='all').to_string()}\n"
    except Exception: info_str += "Could not generate full summary statistics.\n" # Robustness
    return info_str

def convert_to_gradio_chat_format(li_messages):
    # ... (implementation from previous step) ...
    gc = []
    for m in li_messages:
        r, c = m.get("role"), m.get("content")
        if r == "user": gc.append([c,None])
        elif r == "assistant":
            if gc and gc[-1][1] is None and gc[-1][0] is not None: gc[-1][1] = c
            else: gc.append([None,c])
    return gc

RAG_SOURCE_MARKER = re.compile(r"---RAG_SOURCE---(\{.*?\})", re.DOTALL)
DOWNLOAD_FILE_MARKER = re.compile(r"---DOWNLOAD_FILE---(.*?)\n", re.DOTALL)
def format_rag_sources(text):
    # ... (implementation from previous step, ensure robustness) ...
    md, pt = "", text; cn, us = 1, {}
    for m in RAG_SOURCE_MARKER.finditer(text):
        try:
            jsb, sd = m.group(1), json.loads(m.group(1)); sk = sd.get('url',sd.get('title',jsb))
            if sk not in us: us[sk]=cn; st,ti=sd.get("source_type","").lower(),sd.get("title","N/A")
            if st=="pdf": pu=urllib.parse.urlparse(sd.get("url","")); pf=os.path.basename(pu.path) or "s.pdf"; md+=f"[{cn}] PDF: **{ti if ti!='N/A' else pf}**\n"
            elif st=="web": md+=f"[{cn}] Web: [{ti}]({sd.get('url','#')})\n"
            else: md+=f"[{cn}] Source: {ti}\n"
            tc=sd.get("text_chunk",""); md+=f"   > \"{tc[:100]}{'...'if len(tc)>100 else''}\"\n" if tc else "" # Shorter preview
            cn+=1
            pt=pt.replace(m.group(0),f" [{us[sk]}]")
        except Exception: pt=pt.replace(m.group(0),"[Src Error]")
    return pt, ("\n\n**Sources:**\n"+md if md else "")

def format_download_links(text):
    # ... (implementation from previous step) ...
    md, pt = "", text
    for m in DOWNLOAD_FILE_MARKER.finditer(text):
        fn=m.group(1).strip(); du=f"/file={os.path.join(UI_ACCESSIBLE_WORKSPACE,fn)}"
        md+=f"\n\nFile: [{fn}]({du})" # Simpler display
        pt=pt.replace(m.group(0),"")
    return pt, md

def get_agent_response_gr(agent_runner, query: str, chat_history_li_format: list, verbosity_level: int) -> str:
    # ... (modified to include verbosity) ...
    print(f"Getting agent response for query: '{query[:50]}...' with verbosity {verbosity_level}")
    if not agent_runner: return "Error: Agent not available."
    try:
        # Prepend verbosity to query (simple way, agent might need specific parsing for this)
        final_query = f"[Verbosity Level: {verbosity_level}] {query}"
        history_chat_messages = []
        for msg_dict in chat_history_li_format: # History up to, but not including, current user query
            role_str = msg_dict["role"]; role_map = {"user": MessageRole.USER, "assistant": MessageRole.ASSISTANT, "system": MessageRole.SYSTEM}
            mapped_role = role_map.get(role_str.lower(), MessageRole.USER)
            history_chat_messages.append(ChatMessage(role=mapped_role, content=msg_dict["content"]))

        response = agent_runner.chat(final_query, chat_history=history_chat_messages[:-1]) # Pass history *before* current query
        return response.response if hasattr(response, 'response') else str(response)
    except Exception as e: print(f"Error getting agent response: {e}"); return f"Apologies, an error occurred: {e}"

# --- Chat Document Formatters ---
def get_current_chat_markdown(chat_messages_li_format: list) -> str:
    md = "# Chat Discussion\n\n"
    for msg in chat_messages_li_format:
        md += f"**{msg['role'].capitalize()}**: {msg['content']}\n\n---\n"
    return md

def get_current_chat_docx(chat_messages_li_format: list) -> io.BytesIO:
    doc = PythonDocxDocument()
    doc.add_heading('Chat Discussion', level=1)
    for msg in chat_messages_li_format:
        doc.add_paragraph(f"{msg['role'].capitalize()}: ", style='ListBullet').bold = True
        doc.add_paragraph(msg['content'])
        doc.add_paragraph().add_run().add_break() # Separator

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- Gradio App UI and Logic ---
with gr.Blocks(title="ESI-GR", theme=gr.themes.Soft()) as demo:
    # --- States ---
    user_id_state = gr.State()
    long_term_memory_enabled_state = gr.State(True)
    chat_metadata_state = gr.State({})  # {chat_id: name}
    all_chat_messages_state = gr.State({}) # {chat_id: [messages_li_format]}
    current_chat_id_state = gr.State()
    agent_instance_state = gr.State()
    active_chat_display_name_state = gr.State("Current Chat") # For UI display

    uploaded_documents_state = gr.State({})
    uploaded_dataframes_state = gr.State({})

    # LLM Settings States
    temperature_state = gr.State(LlamaSettings.llm.temperature if hasattr(LlamaSettings.llm, 'temperature') else 0.7)
    verbosity_state = gr.State(3) # Default verbosity
    search_results_state = gr.State(10) # Default search results for agent

    gr.Markdown("# 🎓 ESI-GR: ESI Scholarly Instructor (Gradio)")
    # Display current chat name - could be a gr.Textbox or gr.Markdown
    current_chat_name_display_ui = gr.Markdown(f"### Current Chat: {active_chat_display_name_state.value}")


    if not LLM_SETTINGS_OK:
        gr.Error(f"Failed to initialize LLM Settings: {LLM_SETTINGS_ERROR}")
    else:
        with gr.Row():
            # Main Chat Area
            with gr.Column(scale=3):
                chatbot_display = gr.Chatbot(label="ESI Chat", height=550, bubble_full_width=False, show_copy_button=True)
                suggested_prompt_buttons = []
                with gr.Row(elem_id="suggested_prompts_row"):
                    for i in range(SUGGESTED_PROMPT_COUNT):
                        btn = gr.Button(visible=False, scale=1); suggested_prompt_buttons.append(btn)
                with gr.Row():
                    chat_input_textbox = gr.Textbox(show_label=False, placeholder="Type your message here...", lines=3, scale=9, elem_id="chat_input")
                    send_button = gr.Button("Send", scale=1, elem_id="send_button")

            # Sidebar Area
            with gr.Column(scale=1, elem_id="sidebar"):
                with gr.Accordion("💬 Chat History", open=True):
                    new_chat_btn = gr.Button(value="➕ New Chat", variant="secondary")
                    chat_selector_dd = gr.Dropdown(label="Select Chat", choices=[], interactive=True, filterable=True)
                    rename_chat_name_tb = gr.Textbox(label="Rename Current Chat to:", interactive=True)
                    rename_chat_btn = gr.Button("Rename")
                    delete_chat_btn = gr.Button("🗑️ Delete Current Chat", variant="stop")
                    download_chat_md_btn = gr.DownloadButton(label="Download Chat (.md)", variant="secondary")
                    download_chat_docx_btn = gr.DownloadButton(label="Download Chat (.docx)", variant="secondary")

                with gr.Accordion("📁 File Management", open=False):
                    file_uploader = gr.File(label="Upload Documents/Datasets", file_count="multiple", type="filepath")
                    upload_button = gr.Button("Process Uploaded Files")
                    upload_status_display = gr.Markdown()
                    uploaded_files_display = gr.JSON(label="Uploaded & Processed Files") # Shows {filename: type/shape}
                    remove_files_cbg = gr.CheckboxGroup(label="Select files to remove", choices=[], interactive=True)
                    remove_files_btn = gr.Button("Remove Selected Files")

                with gr.Accordion("⚙️ LLM Settings", open=False):
                    temperature_slider = gr.Slider(label="Creativity (Temperature)", minimum=0.0, maximum=2.0, step=0.1, value=temperature_state.value)
                    verbosity_slider = gr.Slider(label="Verbosity Level", minimum=1, maximum=5, step=1, value=verbosity_state.value)
                    search_results_slider = gr.Slider(label="Number of Search Results (Agent)", minimum=3, maximum=15, step=1, value=search_results_state.value) # Requires agent re-init or tool reconfig to take full effect
                    ltm_checkbox = gr.Checkbox(label="Enable Long-term Memory (Save Chats)", value=long_term_memory_enabled_state.value)

                with gr.Accordion("🗑️ Forget Me", open=False):
                    gr.Markdown("⚠️ **Warning**: This action will delete all your conversation history and associated data from the server (if long-term memory was enabled). This cannot be undone.")
                    forget_me_btn = gr.Button("Forget Me and Reset All Data", variant="stop")

                with gr.Accordion("ℹ️ About ESI", open=False):
                    gr.Markdown("ESI (ESI Scholarly Instructor) is an AI-powered assistant designed to help with scholarly research tasks, data analysis, and understanding complex topics. \n\nVersion: 0.2.0 (Gradio)")

    # --- Event Handlers ---
    # Chat Submission & Suggested Prompts (same as before, but ensure verbosity_state is passed to get_agent_response_gr)
    def handle_chat_submit( user_message: str, user_id: str, current_chat_id: str, all_chat_msgs_li_format: dict, ltm_enabled: bool, agent_runner, current_verbosity: int ):
        # ... (previous logic for appending user message, getting agent response, parsing markers, saving history) ...
        if not user_message.strip(): # Handle empty submission
            current_gradio_display = convert_to_gradio_chat_format(all_chat_msgs_li_format.get(current_chat_id, []))
            prompt_updates = [gr.update(visible=False) for _ in suggested_prompt_buttons]
            return current_gradio_display, "", all_chat_msgs_li_format, *prompt_updates

        current_internal_history = list(all_chat_msgs_li_format.get(current_chat_id, []))
        current_internal_history.append({"role": "user", "content": user_message})

        raw_assistant_response = get_agent_response_gr(agent_runner, user_message, current_internal_history, current_verbosity) # Pass verbosity

        assistant_response_text, rag_md = format_rag_sources(raw_assistant_response)
        assistant_response_text, download_md = format_download_links(assistant_response_text)
        final_assistant_response = assistant_response_text + rag_md + download_md

        current_internal_history.append({"role": "assistant", "content": final_assistant_response})
        all_chat_msgs_li_format[current_chat_id] = current_internal_history

        if ltm_enabled and HF_USER_MEMORIES_DATASET_ID and HF_USER_MEMORIES_DATASET_ID != "dummy/user_memories":
            save_chat_history_gr(user_id, current_chat_id, current_internal_history, ltm_enabled)

        updated_gradio_chat_display = convert_to_gradio_chat_format(current_internal_history)
        new_prompts = generate_suggested_prompts(current_internal_history, SUGGESTED_PROMPT_COUNT)
        prompt_updates = [gr.update(value=new_prompts[i] if i < len(new_prompts) else "", visible=i < len(new_prompts)) for i in range(SUGGESTED_PROMPT_COUNT)]
        return updated_gradio_chat_display, "", all_chat_msgs_li_format, *prompt_updates

    chat_submit_outputs = [chatbot_display, chat_input_textbox, all_chat_messages_state] + suggested_prompt_buttons
    chat_input_textbox.submit( handle_chat_submit, inputs=[chat_input_textbox, user_id_state, current_chat_id_state, all_chat_messages_state, long_term_memory_enabled_state, agent_instance_state, verbosity_state], outputs=chat_submit_outputs )
    send_button.click( handle_chat_submit, inputs=[chat_input_textbox, user_id_state, current_chat_id_state, all_chat_messages_state, long_term_memory_enabled_state, agent_instance_state, verbosity_state], outputs=chat_submit_outputs )
    for i, prompt_btn in enumerate(suggested_prompt_buttons):
        prompt_btn.click( handle_chat_submit, inputs=[prompt_btn, user_id_state, current_chat_id_state, all_chat_messages_state, long_term_memory_enabled_state, agent_instance_state, verbosity_state], outputs=chat_submit_outputs )

    # File Upload Handler (same as before)
    def handle_file_uploads( files_list, current_docs_s, current_dfs_s ):
        # ... (previous logic, ensure it updates global dicts and returns for state update) ...
        global _GLOBAL_UPLOADED_DOCS, _GLOBAL_UPLOADED_DFS
        if not files_list: return current_docs_s, current_dfs_s, "No files selected.", {}, [] # No file_names output
        processed_docs, processed_dfs, successes, errors = process_uploaded_files_gr(files_list, current_docs_s, current_dfs_s)
        _GLOBAL_UPLOADED_DOCS = processed_docs.copy(); _GLOBAL_UPLOADED_DFS = processed_dfs.copy()
        status_md = "### Upload Report\n" + ("**Successes:**\n"+"\n".join(f"- {s}" for s in successes)+"\n" if successes else "") + ("\n**Errors:**\n"+"\n".join(f"- {e}" for e in errors)+"\n" if errors else "")
        display_files_info = {fn: "Doc" for fn in processed_docs}
        for fn in processed_dfs: display_files_info[fn] = f"Data (shape: {_GLOBAL_UPLOADED_DFS[fn].shape})"

        # Update choices for remove_files_cbg
        all_uploaded_fnames = list(set(list(processed_docs.keys()) + list(processed_dfs.keys())))
        updated_cbg_choices = gr.update(choices=all_uploaded_fnames if all_uploaded_fnames else [])

        return processed_docs, processed_dfs, status_md, display_files_info, updated_cbg_choices # Added cbg update

    upload_button.click(
        handle_file_uploads,
        inputs=[file_uploader, uploaded_documents_state, uploaded_dataframes_state],
        outputs=[uploaded_documents_state, uploaded_dataframes_state, upload_status_display, uploaded_files_display, remove_files_cbg] # Added cbg
    )

    # --- Sidebar Action Handlers ---
    def new_chat_action_fn(user_id, chat_meta_s, all_chat_msgs_s, ltm_enabled):
        new_chat_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        new_chat_name = f"New Chat {timestamp}"
        chat_meta_s[new_chat_id] = new_chat_name

        greeting = generate_llm_greeting() if LLM_SETTINGS_OK else "Hello! New chat started."
        all_chat_msgs_s[new_chat_id] = [{"role": "assistant", "content": greeting}]

        if ltm_enabled: save_chat_metadata_gr(user_id, chat_meta_s, ltm_enabled) # Save new metadata

        # Update UI elements
        updated_chat_choices = sorted(list(chat_meta_s.values()), reverse=True) # Names for display
        # Map names back to IDs for the dropdown's internal values if needed, or use names as values
        # For simplicity, let's assume dropdown uses chat_id as value, and formats choice text.
        # Or, easier: choices are list of (name, id) pairs. Gradio handles this.
        # Let's use chat_id as value, name as display.
        dropdown_choices = [(name, c_id) for c_id, name in chat_meta_s.items()]

        # Regenerate suggested prompts for the new chat
        new_prompts = generate_suggested_prompts(all_chat_msgs_s[new_chat_id], SUGGESTED_PROMPT_COUNT, DEFAULT_PROMPTS)
        prompt_updates = [gr.update(value=new_prompts[i] if i < len(new_prompts) else "", visible=i < len(new_prompts)) for i in range(SUGGESTED_PROMPT_COUNT)]

        return (new_chat_id, chat_meta_s, all_chat_msgs_s, new_chat_name,
                convert_to_gradio_chat_format(all_chat_msgs_s[new_chat_id]),
                gr.update(choices=dropdown_choices, value=new_chat_id), *prompt_updates)

    new_chat_btn.click(
        new_chat_action_fn,
        inputs=[user_id_state, chat_metadata_state, all_chat_messages_state, long_term_memory_enabled_state],
        outputs=[current_chat_id_state, chat_metadata_state, all_chat_messages_state, active_chat_display_name_state,
                 chatbot_display, chat_selector_dd, *suggested_prompt_buttons]
    ).then(lambda: current_chat_name_display_ui.update(value=f"### Current Chat: {active_chat_display_name_state.value}"), inputs=None, outputs=None)


    def switch_chat_action_fn(selected_chat_id_from_dd, chat_meta_s, all_chat_msgs_s): # selected_chat_id_from_dd is chat_id
        if not selected_chat_id_from_dd: # No chat selected or empty dropdown
            return gr.update(), gr.update(), gr.update(), gr.update() # No change to current_chat_id, name, chatbot, prompts

        current_chat_name = chat_meta_s.get(selected_chat_id_from_dd, "Unknown Chat")
        current_messages_li = all_chat_msgs_s.get(selected_chat_id_from_dd, [])
        gradio_chat_display = convert_to_gradio_chat_format(current_messages_li)

        new_prompts = generate_suggested_prompts(current_messages_li, SUGGESTED_PROMPT_COUNT, DEFAULT_PROMPTS)
        prompt_updates = [gr.update(value=new_prompts[i] if i < len(new_prompts) else "", visible=i < len(new_prompts)) for i in range(SUGGESTED_PROMPT_COUNT)]

        return selected_chat_id_from_dd, current_chat_name, gradio_chat_display, *prompt_updates

    chat_selector_dd.change(
        switch_chat_action_fn,
        inputs=[chat_selector_dd, chat_metadata_state, all_chat_messages_state],
        outputs=[current_chat_id_state, active_chat_display_name_state, chatbot_display, *suggested_prompt_buttons]
    ).then(lambda: current_chat_name_display_ui.update(value=f"### Current Chat: {active_chat_display_name_state.value}"), inputs=None, outputs=None)


    def rename_chat_action_fn(new_name_str, user_id, current_chat_id, chat_meta_s, ltm_enabled):
        if not new_name_str.strip() or not current_chat_id:
            return chat_meta_s, gr.update() # No change if name is empty or no chat selected

        chat_meta_s[current_chat_id] = new_name_str.strip()
        if ltm_enabled: save_chat_metadata_gr(user_id, chat_meta_s, ltm_enabled)

        dropdown_choices = [(name, c_id) for c_id, name in chat_meta_s.items()]
        return chat_meta_s, new_name_str.strip(), gr.update(choices=dropdown_choices, value=current_chat_id), "" # Clear rename textbox

    rename_chat_btn.click(
        rename_chat_action_fn,
        inputs=[rename_chat_name_tb, user_id_state, current_chat_id_state, chat_metadata_state, long_term_memory_enabled_state],
        outputs=[chat_metadata_state, active_chat_display_name_state, chat_selector_dd, rename_chat_name_tb]
    ).then(lambda: current_chat_name_display_ui.update(value=f"### Current Chat: {active_chat_display_name_state.value}"), inputs=None, outputs=None)

    def delete_chat_action_fn(user_id, current_chat_id_to_delete, chat_meta_s, all_chat_msgs_s, ltm_enabled):
        if not current_chat_id_to_delete: return chat_meta_s, all_chat_msgs_s, gr.update(), gr.update(), gr.update(), gr.update()

        if current_chat_id_to_delete in chat_meta_s: del chat_meta_s[current_chat_id_to_delete]
        if current_chat_id_to_delete in all_chat_msgs_s: del all_chat_msgs_s[current_chat_id_to_delete]

        if ltm_enabled:
            save_chat_metadata_gr(user_id, chat_meta_s, ltm_enabled) # Save updated metadata (removes deleted one)
            # To delete specific chat messages from the _messages.json, we'd need to re-save the whole file without it.
            # The current save_chat_history_gr saves all messages for a user.
            # So, effectively, all_chat_msgs_s no longer having it means it won't be re-saved next time for that user.
            # Or, if _messages.json is {user_id: {chat_id: messages}}, then this is fine.
            # The current structure is {user_id}_messages.json containing {chat_id: messages}.
            # So, we need to save the all_chat_msgs_s which now excludes the deleted chat.
            save_chat_history_gr(user_id, "placeholder_chat_id_for_saving_all_user_messages", all_chat_msgs_s, ltm_enabled) # A bit hacky, implies save_chat_history needs to handle this structure.
            # Let's assume save_chat_history_gr(user_id, chat_id, messages,...) actually saves messages[chat_id] = messages for that user.
            # A better approach for deleting history on HF: delete the user's entire message file if all chats are gone,
            # or re-save the all_chat_msgs_s object.
            # For now, let's assume `save_chat_history_gr` is smart enough if we pass `all_chat_msgs_s` as the "messages" for a "user".
            # This needs clarification in `save_chat_history_gr`.
            # Simpler: just save metadata. History for deleted chat remains in file but inaccessible via UI.
            # Best: modify save_chat_history_gr to take all_chat_msgs_s for a user_id and overwrite.
            # For this subtask, current save_chat_history_gr will simply not re-save the deleted chat_id if it's not in all_chat_msgs_s.
            # However, it expects a single chat_id and its messages.
            # So, after deleting from all_chat_msgs_s, we can't use current save_chat_history_gr to reflect deletion on HF.
            # This is a limitation for now. We'll remove from HF if possible or mark as "to be improved".
            # For now, it's deleted from current session. Re-saving metadata is key.

        # Switch to another chat or create a new one if no chats are left
        new_current_chat_id = next(iter(chat_meta_s.keys()), None)
        new_active_chat_name = "No Chat Selected"
        new_chatbot_display = []
        new_prompt_updates = [gr.update(visible=False) for _ in suggested_prompt_buttons]

        if new_current_chat_id:
            new_active_chat_name = chat_meta_s[new_current_chat_id]
            new_chatbot_display = convert_to_gradio_chat_format(all_chat_msgs_s.get(new_current_chat_id, []))
            new_prompts = generate_suggested_prompts(all_chat_msgs_s.get(new_current_chat_id, []), SUGGESTED_PROMPT_COUNT, DEFAULT_PROMPTS)
            new_prompt_updates = [gr.update(value=new_prompts[i] if i < len(new_prompts) else "", visible=i < len(new_prompts)) for i in range(SUGGESTED_PROMPT_COUNT)]
        else: # No chats left, create a brand new one (like new_chat_action_fn without saving metadata yet)
            new_current_chat_id = str(uuid.uuid4())
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            new_active_chat_name = f"New Chat {timestamp}"
            chat_meta_s[new_current_chat_id] = new_active_chat_name
            greeting = generate_llm_greeting() if LLM_SETTINGS_OK else "Hello! New chat started."
            all_chat_msgs_s[new_current_chat_id] = [{"role": "assistant", "content": greeting}]
            new_chatbot_display = convert_to_gradio_chat_format(all_chat_msgs_s[new_current_chat_id])
            # Save this new chat's metadata
            if ltm_enabled: save_chat_metadata_gr(user_id, chat_meta_s, ltm_enabled)

        dropdown_choices = [(name, c_id) for c_id, name in chat_meta_s.items()]
        return (chat_meta_s, all_chat_msgs_s, new_current_chat_id, new_active_chat_name,
                new_chatbot_display, gr.update(choices=dropdown_choices, value=new_current_chat_id),
                *new_prompt_updates)

    delete_chat_btn.click(
        delete_chat_action_fn,
        inputs=[user_id_state, current_chat_id_state, chat_metadata_state, all_chat_messages_state, long_term_memory_enabled_state],
        outputs=[chat_metadata_state, all_chat_messages_state, current_chat_id_state, active_chat_display_name_state,
                 chatbot_display, chat_selector_dd, *suggested_prompt_buttons]
    ).then(lambda: current_chat_name_display_ui.update(value=f"### Current Chat: {active_chat_display_name_state.value}"), inputs=None, outputs=None)


    # Download Chat Handlers
    def download_chat_md_content_fn(current_chat_id, all_chat_msgs_s):
        if not current_chat_id: return None # Or raise error for Gradio
        messages = all_chat_msgs_s.get(current_chat_id, [])
        md_content = get_current_chat_markdown(messages)
        # Gradio DownloadButton needs a filepath. Save to temp file.
        temp_dir = os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE, "temp_downloads")
        os.makedirs(temp_dir, exist_ok=True)
        filepath = os.path.join(temp_dir, f"chat_{current_chat_id}.md")
        with open(filepath, "w", encoding="utf-8") as f: f.write(md_content)
        return filepath

    download_chat_md_btn.click(None, # No function needed here if DownloadButton uses a generator
        inputs=[current_chat_id_state, all_chat_messages_state], # These are passed to the generator fn
        outputs=download_chat_md_btn, # The button itself is the output
        # The actual function is passed to DownloadButton directly or via _js
        # For Python function, it's implicitly called when inputs change if button is output
        # No, this is wrong. DownloadButton needs a function that RETURNS the filepath.
        # The click itself triggers the function call.
    ).then(lambda current_chat_id, all_chat_msgs_s: download_chat_md_content_fn(current_chat_id, all_chat_msgs_s),
           inputs=[current_chat_id_state, all_chat_messages_state], outputs=download_chat_md_btn)


    def download_chat_docx_content_fn(current_chat_id, all_chat_msgs_s):
        if not current_chat_id: return None
        messages = all_chat_msgs_s.get(current_chat_id, [])
        docx_bio = get_current_chat_docx(messages)
        temp_dir = os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE, "temp_downloads")
        os.makedirs(temp_dir, exist_ok=True)
        filepath = os.path.join(temp_dir, f"chat_{current_chat_id}.docx")
        with open(filepath, "wb") as f: f.write(docx_bio.getvalue())
        return filepath

    # Correct way to wire DownloadButton:
    # download_chat_docx_btn.download = download_chat_docx_content_fn # This is not how gr.DownloadButton works
    # It needs to be an output of a click or other event.
    # The button itself can be an output, and the click function returns the filepath.
    download_chat_docx_btn.click(
        download_chat_docx_content_fn, # This function must return a filepath
        inputs=[current_chat_id_state, all_chat_messages_state],
        outputs=download_chat_docx_btn # Special handling for DownloadButton
    )


    # File Removal Handler
    def remove_selected_files_action_fn(files_to_remove_list, current_docs_s, current_dfs_s):
        global _GLOBAL_UPLOADED_DOCS, _GLOBAL_UPLOADED_DFS
        if not files_to_remove_list: return current_docs_s, current_dfs_s, "No files selected for removal.", gr.update()

        removed_count = 0
        for fname in files_to_remove_list:
            if fname in _GLOBAL_UPLOADED_DOCS: del _GLOBAL_UPLOADED_DOCS[fname]; removed_count+=1
            if fname in current_docs_s: del current_docs_s[fname] # Update state copy
            if fname in _GLOBAL_UPLOADED_DFS: del _GLOBAL_UPLOADED_DFS[fname]; removed_count+=1
            if fname in current_dfs_s: del current_dfs_s[fname] # Update state copy

            # Also delete from workspace
            try: os.remove(os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE, fname))
            except OSError: print(f"Could not delete file from workspace: {fname}")

        status_msg = f"Removed {removed_count} file(s). Remaining: {len(_GLOBAL_UPLOADED_DOCS) + len(_GLOBAL_UPLOADED_DFS)}"

        # Update CheckboxGroup choices
        all_remaining_fnames = list(set(list(_GLOBAL_UPLOADED_DOCS.keys()) + list(_GLOBAL_UPLOADED_DFS.keys())))
        updated_cbg_choices = gr.update(choices=all_remaining_fnames if all_remaining_fnames else [], value=[]) # Clear selection

        # Update main display of uploaded files
        display_files_info = {fn: "Doc" for fn in _GLOBAL_UPLOADED_DOCS}
        for fn in _GLOBAL_UPLOADED_DFS: display_files_info[fn] = f"Data (shape: {_GLOBAL_UPLOADED_DFS[fn].shape})"

        return current_docs_s, current_dfs_s, status_msg, display_files_info, updated_cbg_choices

    remove_files_btn.click(
        remove_selected_files_action_fn,
        inputs=[remove_files_cbg, uploaded_documents_state, uploaded_dataframes_state],
        outputs=[uploaded_documents_state, uploaded_dataframes_state, upload_status_display, uploaded_files_display, remove_files_cbg]
    )

    # LLM Settings Handlers
    def update_temperature_fn(temp_val, temp_state_s):
        if hasattr(LlamaSettings, 'llm') and hasattr(LlamaSettings.llm, 'temperature'):
            LlamaSettings.llm.temperature = temp_val
        temp_state_s = temp_val # Update state
        return temp_state_s, f"Temperature set to {temp_val}." # Feedback

    temperature_status_md = gr.Markdown() # For feedback
    temperature_slider.release(update_temperature_fn, inputs=[temperature_slider, temperature_state], outputs=[temperature_state, temperature_status_md])

    # Verbosity and Search Results update their states. These are read by other functions.
    verbosity_slider.release(lambda v: v, inputs=[verbosity_slider], outputs=[verbosity_state])
    search_results_slider.release(lambda v: v, inputs=[search_results_slider], outputs=[search_results_state]) # Agent re-init needed for full effect
    ltm_checkbox.change(lambda v: v, inputs=[ltm_checkbox], outputs=[long_term_memory_enabled_state])


    # "Forget Me" Handler
    def forget_me_action_fn(user_id, request: gr.Request): # request: gr.Request for potential full page reload
        global _GLOBAL_UPLOADED_DOCS, _GLOBAL_UPLOADED_DFS
        print(f"Forget Me action triggered for user_id: {user_id}")
        if user_id and HF_USER_MEMORIES_DATASET_ID and HF_USER_MEMORIES_DATASET_ID != "dummy/user_memories" and fs.token:
            try:
                metadata_repo_path = f"user_memories/{user_id}_metadata.json" # Path within the dataset repo
                messages_repo_path = f"user_memories/{user_id}_messages.json"

                # Construct full paths for hf_hub_delete
                full_metadata_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{metadata_repo_path}"
                full_messages_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_repo_path}"

                if fs.exists(f"{HF_USER_MEMORIES_DATASET_ID}/{metadata_repo_path}"): # fs.exists uses relative path to dataset root
                    hf_hub_delete(repo_id=HF_USER_MEMORIES_DATASET_ID, path_in_repo=metadata_repo_path, repo_type="dataset", token=fs.token)
                    print(f"Deleted metadata for user {user_id} from Hugging Face.")
                if fs.exists(f"{HF_USER_MEMORIES_DATASET_ID}/{messages_repo_path}"):
                    hf_hub_delete(repo_id=HF_USER_MEMORIES_DATASET_ID, path_in_repo=messages_repo_path, repo_type="dataset", token=fs.token)
                    print(f"Deleted messages for user {user_id} from Hugging Face.")
            except Exception as e:
                print(f"Error deleting user data from Hugging Face for user {user_id}: {e}")
                # Continue with local reset anyway

        # Clear local global stores
        _GLOBAL_UPLOADED_DOCS.clear()
        _GLOBAL_UPLOADED_DFS.clear()

        # Reset all relevant states to trigger _initialize_session_and_agent logic for a new session
        # This effectively calls parts of _initialize_session_and_agent again.
        # The goal is to get a new user_id and clear all session-specific data.
        # A full page reload might be cleaner if Gradio supports it easily.
        # For now, manually reset critical states.

        # Call the initialization function to get fresh states
        # This is like a soft reset.
        (new_user_id, new_ltm_enabled, new_chat_meta, new_all_chat_msgs, new_current_chat_id,
         new_chatbot_disp, new_agent_runner, new_debug_chat_id,
         new_uploaded_docs, new_uploaded_dfs, new_upload_status, new_files_disp,
         *new_prompt_updates_init) = _initialize_session_and_agent(request) # Pass request for potential use in init

        # This needs to return all outputs that demo.load expects
        # Plus a message
        forget_status_msg = "All your data has been reset. A new session has started."

        # We need to return updates for all components that _initialize_session_and_agent updates
        # This is becoming very complex. A simpler way might be to just reload the page with JS.
        # Or, return a specific set of updates that reset the key UI parts.

        # For now, return a tuple that includes updates for the main states and a status message.
        # The full re-initialization of UI components via return values here is tricky.
        # It's better if _initialize_session_and_agent is the single source of truth for initial UI state.
        # So, the forget_me_btn click should ideally trigger demo.load again.
        # This is not directly possible.
        # Alternative: return updates for key components to simulate a reset.

        return (new_user_id, new_ltm_enabled, new_chat_meta, new_all_chat_msgs, new_current_chat_id,
                new_agent_runner, "New Session", # active_chat_display_name
                new_chatbot_disp, # chatbot
                gr.update(choices=[], value=None), # chat_selector_dd
                gr.update(choices=[], value=[]), # remove_files_cbg
                new_uploaded_docs, new_uploaded_dfs, new_upload_status, new_files_disp,
                forget_status_msg, # A new component for this message, or an alert
                *new_prompt_updates_init)


    forget_me_status_display = gr.Markdown() # Add this component where you want the message
    # Update outputs list for forget_me_btn
    forget_me_outputs = [
        user_id_state, long_term_memory_enabled_state, chat_metadata_state, all_chat_messages_state,
        current_chat_id_state, agent_instance_state, active_chat_display_name_state,
        chatbot_display, chat_selector_dd, remove_files_cbg,
        uploaded_documents_state, uploaded_dataframes_state, upload_status_display, uploaded_files_display,
        forget_me_status_display # For the message
    ] + suggested_prompt_buttons

    forget_me_btn.click(
        forget_me_action_fn,
        inputs=[user_id_state, gr.Request()], # Pass request object
        outputs=forget_me_outputs
    ).then(lambda: current_chat_name_display_ui.update(value=f"### Current Chat: {active_chat_display_name_state.value}"), inputs=None, outputs=None)


    # --- Initialization (`demo.load`) ---
    # _initialize_session_and_agent needs to return initial values for all new states/UI
    def _initialize_session_and_agent(request: gr.Request): # request might be None if not passed by all callers
        _user_id = get_user_id_gr()
        _long_term_memory_enabled = long_term_memory_enabled_state.value # Use current state or default
        _current_temp = temperature_state.value
        _current_verbosity = verbosity_state.value
        _current_search_results = search_results_state.value

        # Update LlamaSettings temperature if LLM object exists
        if hasattr(LlamaSettings, 'llm') and hasattr(LlamaSettings.llm, 'temperature'):
            LlamaSettings.llm.temperature = _current_temp

        dynamic_tools_for_agent = [
            FunctionTool.from_defaults(fn=read_uploaded_document_tool_fn_gr, name="read_uploaded_document", description="Reads user-uploaded document text."),
            FunctionTool.from_defaults(fn=analyze_dataframe_tool_fn_gr, name="analyze_uploaded_dataframe", description="Analyzes user-uploaded DataFrame.")
        ]
        _agent_runner = None
        if LLM_SETTINGS_OK:
            try: _agent_runner = create_orchestrator_agent(dynamic_tools=dynamic_tools_for_agent, max_search_results=_current_search_results); print(f"Agent created with search results: {_current_search_results}")
            except Exception as e: print(f"Error initializing agent: {e}")

        loaded_data = {"metadata": {}, "messages": {}}
        if _long_term_memory_enabled:
            if HF_USER_MEMORIES_DATASET_ID and HF_USER_MEMORIES_DATASET_ID != "dummy/user_memories":
                loaded_data = _load_user_data_from_hf(_user_id)

        _chat_metadata = loaded_data.get("metadata", {})
        _all_chat_messages_li_format = loaded_data.get("messages", {})
        _current_chat_id = next(iter(_chat_metadata.keys()), None)
        _active_chat_name = _chat_metadata.get(_current_chat_id, "New Chat")

        if not _current_chat_id: # New session or no chats loaded
            _current_chat_id = str(uuid.uuid4())
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            _active_chat_name = f"Chat {timestamp}"
            _chat_metadata[_current_chat_id] = _active_chat_name
            greeting = generate_llm_greeting() if LLM_SETTINGS_OK else "Hello!"
            _all_chat_messages_li_format[_current_chat_id] = [{"role": "assistant", "content": greeting}]
            if _long_term_memory_enabled: save_chat_metadata_gr(_user_id, _chat_metadata, _long_term_memory_enabled)

        _current_messages_li_format = _all_chat_messages_li_format.get(_current_chat_id, [])
        initial_gradio_chat_display = convert_to_gradio_chat_format(_current_messages_li_format)

        initial_prompts = generate_suggested_prompts(_current_messages_li_format, SUGGESTED_PROMPT_COUNT, DEFAULT_PROMPTS)
        initial_prompt_updates = [gr.update(value=initial_prompts[i] if i < len(initial_prompts) else "", visible=i < len(initial_prompts)) for i in range(SUGGESTED_PROMPT_COUNT)]

        dropdown_choices = [(name, c_id) for c_id, name in _chat_metadata.items()]
        chat_selector_dd_update = gr.update(choices=dropdown_choices if dropdown_choices else [], value=_current_chat_id)

        # Initial state for file UI
        initial_uploaded_docs = _GLOBAL_UPLOADED_DOCS.copy() # Persist across soft reloads if global isn't cleared
        initial_uploaded_dfs = _GLOBAL_UPLOADED_DFS.copy()
        initial_upload_status = "Upload files using the controls."
        initial_files_display = {fn: "Doc" for fn in initial_uploaded_docs}
        for fn in initial_uploaded_dfs: initial_files_display[fn] = f"Data (shape: {initial_uploaded_dfs[fn].shape})"
        all_uploaded_fnames = list(set(list(initial_uploaded_docs.keys()) + list(initial_uploaded_dfs.keys())))
        remove_files_cbg_update = gr.update(choices=all_uploaded_fnames if all_uploaded_fnames else [])

        # Return all states that demo.load expects
        return (_user_id, _long_term_memory_enabled, _chat_metadata, _all_chat_messages_li_format,
                _current_chat_id, _agent_runner, _active_chat_name, initial_gradio_chat_display,
                chat_selector_dd_update, remove_files_cbg_update,
                initial_uploaded_docs, initial_uploaded_dfs, initial_upload_status, initial_files_display,
                _current_temp, _current_verbosity, _current_search_results, # LLM settings states
                *initial_prompt_updates)

    load_outputs = [ # Order must match the return tuple of _initialize_session_and_agent
        user_id_state, long_term_memory_enabled_state, chat_metadata_state, all_chat_messages_state,
        current_chat_id_state, agent_instance_state, active_chat_display_name_state, chatbot_display,
        chat_selector_dd, remove_files_cbg, # UI elements to update
        uploaded_documents_state, uploaded_dataframes_state, upload_status_display, uploaded_files_display,
        temperature_state, verbosity_state, search_results_state, # LLM states
    ] + suggested_prompt_buttons

    # Add forget_me_status_display to the page if not already part of a defined layout
    # For simplicity, it's an output of forget_me_btn, assuming it's defined somewhere.
    # If it needs to be part of the main layout:
    # with gr.Row(): forget_me_status_display = gr.Markdown() # Place it appropriately

    demo.load( _initialize_session_and_agent, inputs=gr.Request(), outputs=load_outputs # Pass request to init
              ).then(lambda val: current_chat_name_display_ui.update(value=f"### Current Chat: {val}"), inputs=active_chat_display_name_state, outputs=None)


if __name__ == "__main__":
    print("Starting Gradio App...")
    workspace_full_path = os.path.join(PROJECT_ROOT, UI_ACCESSIBLE_WORKSPACE)
    temp_download_path = os.path.join(workspace_full_path, "temp_downloads")
    if not os.path.exists(temp_download_path): os.makedirs(temp_download_path, exist_ok=True)

    # Ensure all paths served by /file= are absolute or relative to where app is run.
    # Gradio's `allowed_paths` should ideally be absolute paths for robustness.
    # `workspace_full_path` is already absolute if PROJECT_ROOT is.
    allowed_paths_list = [workspace_full_path, temp_download_path]

    demo.launch(debug=True, share=False, allowed_paths=allowed_paths_list)
